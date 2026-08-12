from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
EN_MD = ROOT / "docs" / "why-agent-engineering-needs-controlled-experiments.md"
ZH_MD = ROOT / "docs" / "why-agent-engineering-needs-controlled-experiments.zh-CN.md"
OUT = ROOT / "docs" / "assets" / "agent-engineering-controlled-experiments-bilingual.docx"
EN_FIGURE = ROOT / "docs" / "images" / "six_mismatch.en.png"
ZH_FIGURE = ROOT / "docs" / "images" / "six_mismatch.zh.png"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "0A6E75"
GOLD = "A66C00"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BLACK = "111111"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B8C2CE", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)
        tag.set(qn("w:space"), "0")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths: list[int]) -> None:
    assert sum(widths) == TABLE_WIDTH_DXA
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)


def set_run_font(run, size=None, bold=None, italic=None, color=None, code=False) -> None:
    font_name = "Consolas" if code else "Calibri"
    east_asia = "Microsoft YaHei"
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text: str, url: str, bold=False, italic=False, code=False):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.extend([r_fonts, color, underline])
    if bold:
        r_pr.append(OxmlElement("w:b"))
    if italic:
        r_pr.append(OxmlElement("w:i"))
    if code:
        r_fonts.set(qn("w:ascii"), "Consolas")
        r_fonts.set(qn("w:hAnsi"), "Consolas")
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(r"(\[.+?\]\(https?://[^)]+\)|\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")


def add_inline_markdown(paragraph, text: str, size=11, color=BLACK, bold_default=False) -> None:
    pos = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=size, color=color, bold=bold_default)
        token = match.group(0)
        link_match = re.fullmatch(r"\[(.+?)\]\((https?://[^)]+)\)", token)
        if link_match:
            add_hyperlink(paragraph, link_match.group(1), link_match.group(2), bold=bold_default)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size - 0.3, color=DARK_BLUE, code=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, color=color, italic=True, bold=bold_default)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, color=color, bold=bold_default)


def paragraph_border(paragraph, side="left", color=BLUE, size="12", space="8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), size)
    edge.set(qn("w:space"), space)
    edge.set(qn("w:color"), color)
    p_bdr.append(edge)


def paragraph_shading(paragraph, fill=CALLOUT) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_run_font(run, size=8.5, color=MUTED)


def add_num_definition(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.extend([tabs, ind])
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.append(fonts)
    lvl.extend([start, fmt, lvl_text, suff, p_pr, r_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def set_num_paragraph(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    n_id = OxmlElement("w:numId")
    n_id.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, n_id])


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.widow_control = True

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name, size, color, bold, after in (
        ("Language Title", 24, NAVY, True, 5),
        ("Language Subtitle", 14, DARK_BLUE, False, 16),
        ("Caption Custom", 9, MUTED, False, 8),
        ("Metadata", 9.5, MUTED, False, 4),
    ):
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1


def configure_section(section) -> None:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.different_first_page_header_footer = True


def configure_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    left = p.add_run("AGENT ENGINEERING · CONTROLLED EXPERIMENTS")
    set_run_font(left, size=8, bold=True, color=NAVY)
    p.add_run("\t")
    right = p.add_run("双语研究报告 · BILINGUAL REPORT")
    set_run_font(right, size=8, color=MUTED)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
    paragraph_border(p, side="bottom", color="D7DBE2", size="4", space="3")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("against-llm-mediocrity  ·  2026-08-12  ·  ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_field(fp)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RESEARCH METHODS REPORT · 研究方法报告")
    set_run_font(r, size=10.5, bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Agent 工程为什么必须重视受控实验")
    set_run_font(r, size=29, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run("Why Agent Engineering Must Take Controlled Experiments Seriously")
    set_run_font(r, size=20, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("From benchmark progress to a cumulative science of runtime mechanisms")
    set_run_font(r, size=13.5, italic=True, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.65)
    p.paragraph_format.right_indent = Inches(0.65)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(24)
    p.paragraph_format.line_spacing = 1.3
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_shading(p, "EEF4FA")
    paragraph_border(p, side="left", color=TEAL, size="18", space="12")
    r = p.add_run("模型升级移动失配边界；受控实验识别边界上的治理规律。\n")
    set_run_font(r, size=13, bold=True, color=TEAL)
    r = p.add_run("Model upgrades move mismatch boundaries; controlled experiments identify the governance laws at those boundaries.")
    set_run_font(r, size=11.5, italic=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(50)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Working Draft v0.1 · 2026-08-12")
    set_run_font(r, size=10.5, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("against-llm-mediocrity")
    set_run_font(r, size=9.5, color=MUTED)
    doc.add_page_break()


def add_reading_guide(doc: Document) -> None:
    p = doc.add_paragraph("阅读指南 · Reading Guide", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    add_body_paragraph(doc, "本报告先给出中文全文，再给出英文全文。两种语言采用相同的论证结构、六类实验合同和参考文献集合。")
    add_body_paragraph(doc, "The Chinese text appears first, followed by the English text. Both editions use the same argument structure, six-mismatch experiment contracts, and source set.")

    table = doc.add_table(rows=1, cols=3)
    widths = [1800, 3300, 4260]
    headers = ["部分 / Part", "核心问题 / Core question", "交付 / Deliverable"]
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text
    rows = [
        ("前沿综述", "当前研究已经知道什么？", "证据梯度与来源边界"),
        ("局限性", "为什么 Harness 重要仍不等于因果知识？", "八类归因与评价缺口"),
        ("六类失配", "什么结构条件持续产生失败？", "六套受控实验合同"),
        ("方法合同", "实验怎样才能可复现、可证伪、可迁移？", "最小实验合同与研究路线"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    format_table(table, widths)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    paragraph_shading(p, "FFF8E8")
    paragraph_border(p, side="left", color=GOLD, size="16", space="10")
    add_inline_markdown(p, "证据边界 / Evidence boundary: 2026 年预印本用于描述前沿方向，不被当作已定稿的普遍定律。", size=10.5, color=NAVY, bold_default=True)
    doc.add_page_break()


def add_body_paragraph(doc: Document, text: str, style=None, align=None) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    add_inline_markdown(p, text)


def format_table(table, widths: list[int]) -> None:
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.12
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if paragraph.text:
                    text = paragraph.text
                    paragraph.clear()
                    add_inline_markdown(paragraph, text, size=8.6 if len(widths) >= 4 else 9.0, color=NAVY if r_idx == 0 else BLACK, bold_default=(r_idx == 0))


def table_widths(col_count: int) -> list[int]:
    if col_count == 2:
        return [2700, 6660]
    if col_count == 3:
        return [1800, 3300, 4260]
    if col_count == 4:
        return [1500, 2200, 2700, 2960]
    base = TABLE_WIDTH_DXA // col_count
    widths = [base] * col_count
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return widths


def parse_table(lines: list[str], start: int):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row[:col_count]):
            p = table.cell(r_idx, c_idx).paragraphs[0]
            p.clear()
            add_inline_markdown(p, value, size=8.6 if col_count >= 4 else 9.0, color=NAVY if r_idx == 0 else BLACK, bold_default=(r_idx == 0))
    format_table(table, table_widths(col_count))
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)


def add_code_block(doc: Document, code_lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.05
    paragraph_shading(p, LIGHT_GRAY)
    paragraph_border(p, side="left", color=TEAL, size="14", space="8")
    for idx, line in enumerate(code_lines):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, size=9, color=DARK_BLUE, code=True)


def add_quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.25
    paragraph_shading(p, "EEF4FA")
    paragraph_border(p, side="left", color=BLUE, size="18", space="9")
    add_inline_markdown(p, text, size=11, color=NAVY, bold_default=True)


def add_figure_page(doc: Document, image_path: Path, caption: str) -> None:
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    picture = p.add_run().add_picture(str(image_path), width=Inches(5.08))
    # Keep the figure meaningful for readers who cannot see the bitmap.  Word
    # stores inline-image alternative text on the drawing's docPr element.
    picture._inline.docPr.set("title", "Six primitive mismatches")
    picture._inline.docPr.set("descr", caption)
    cap = doc.add_paragraph(style="Caption Custom")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline_markdown(cap, caption, size=9, color=MUTED)
    doc.add_page_break()


def render_markdown(doc: Document, path: Path, language: str) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    in_math = False
    math_lines: list[str] = []
    active_list_kind = None
    active_num_id = None
    figure_added = False
    abstract_mode = False
    references_mode = False

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(doc, code_lines)
                in_code = False
                code_lines = []
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped == "\\[":
            in_math = True
            math_lines = []
            i += 1
            continue
        if stripped == "\\]" and in_math:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(9)
            for idx, value in enumerate(math_lines):
                if idx:
                    p.add_run().add_break()
                run = p.add_run(value)
                set_run_font(run, size=10.5, color=NAVY)
                run.font.name = "Cambria Math"
                run._element.rPr.rFonts.set(qn("w:ascii"), "Cambria Math")
                run._element.rPr.rFonts.set(qn("w:hAnsi"), "Cambria Math")
            in_math = False
            math_lines = []
            i += 1
            continue
        if in_math:
            math_lines.append(stripped)
            i += 1
            continue

        if not stripped:
            active_list_kind = None
            active_num_id = None
            i += 1
            continue
        if stripped == "---":
            active_list_kind = None
            active_num_id = None
            i += 1
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            hashes, text = heading.groups()
            if text in ("References", "参考文献"):
                doc.add_page_break()
            if text in ("1. The methodological problem is confounding, not a shortage of benchmarks", "1. 方法论问题是混杂，不是 Benchmark 不够多") and not figure_added:
                add_figure_page(
                    doc,
                    ZH_FIGURE if language == "zh" else EN_FIGURE,
                    "图 1. LLM 系统中的六类原始失配与价值保存管线。" if language == "zh" else "Figure 1. Six primitive mismatches across the LLM-system value-preservation pipeline.",
                )
                figure_added = True
            level = len(hashes)
            if level == 1:
                p = doc.add_paragraph(style="Language Title")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_inline_markdown(p, text, size=24, color=NAVY, bold_default=True)
            elif level == 2 and (text.startswith("From ") or text.startswith("从 Benchmark")):
                p = doc.add_paragraph(style="Language Subtitle")
                add_inline_markdown(p, text, size=14, color=DARK_BLUE)
            else:
                p = doc.add_paragraph(style=f"Heading {min(level - 1, 3)}")
                add_inline_markdown(p, text, size={2: 16, 3: 13}.get(level, 12), color=BLUE if level == 2 else DARK_BLUE, bold_default=True)
            abstract_mode = text in ("Abstract", "摘要")
            references_mode = text in ("References", "参考文献")
            if level == 2 and text.startswith(("1. ", "1.")):
                abstract_mode = False
            if level == 2 and text in ("Evidence and claim boundary", "证据与主张边界"):
                references_mode = False
            active_list_kind = None
            active_num_id = None
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows, i = parse_table(lines, i)
            add_markdown_table(doc, rows)
            active_list_kind = None
            active_num_id = None
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet or numbered:
            kind = "bullet" if bullet else "decimal"
            if active_list_kind != kind or active_num_id is None:
                active_num_id = add_num_definition(doc, kind)
                active_list_kind = kind
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2 if references_mode else 4)
            p.paragraph_format.line_spacing = 1.1 if references_mode else 1.208
            set_num_paragraph(p, active_num_id)
            add_inline_markdown(p, (bullet or numbered).group(1), size=9.2 if references_mode else 10.7)
            i += 1
            continue

        if stripped.startswith(">"):
            add_quote(doc, stripped.lstrip("> "))
            active_list_kind = None
            active_num_id = None
            i += 1
            continue

        if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            p = doc.add_paragraph(style="Metadata")
            add_inline_markdown(p, stripped, size=9.5, color=MUTED)
            i += 1
            continue

        p = doc.add_paragraph()
        if abstract_mode:
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.2
            add_inline_markdown(p, stripped, size=10.2)
        else:
            add_inline_markdown(p, stripped)
        active_list_kind = None
        active_num_id = None
        i += 1


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)

    props = doc.core_properties
    props.title = "Agent 工程为什么必须重视受控实验 / Why Agent Engineering Must Take Controlled Experiments Seriously"
    props.subject = "Controlled experiments and six primitive mismatches in Agent Engineering"
    props.author = "against-llm-mediocrity"
    props.keywords = "Agent Engineering, controlled experiments, harness, six primitive mismatches, causal evaluation"
    props.comments = "Bilingual research methods report generated from the repository Markdown sources."

    add_cover(doc)
    add_reading_guide(doc)
    render_markdown(doc, ZH_MD, "zh")
    doc.add_page_break()
    render_markdown(doc, EN_MD, "en")

    for section in doc.sections:
        configure_section(section)
        configure_header_footer(section)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
