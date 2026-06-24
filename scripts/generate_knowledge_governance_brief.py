# -*- coding: utf-8 -*-
from __future__ import annotations

from math import cos, pi, sin
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "tmp" / "assets" / "knowledge_governance_brief"
PDF_PATH = ROOT / "output" / "pdf" / "knowledge-governance-llm-systems-short.zh-CN.pdf"
DOCX_PATH = ROOT / "output" / "documents" / "knowledge-governance-llm-systems-short.zh-CN.docx"

SOURCE_URL = "https://wxy2ab.github.io/against-llm-mediocrity/zh/"

SIMHEI = Path("C:/Windows/Fonts/simhei.ttf")
MSYH = Path("C:/Windows/Fonts/msyh.ttc")
MSYH_BOLD = Path("C:/Windows/Fonts/msyhbd.ttc")

FONT_PATH = SIMHEI if SIMHEI.exists() else MSYH
FONT_BOLD_PATH = MSYH_BOLD if MSYH_BOLD.exists() else FONT_PATH

INK = "#19212A"
MUTED = "#5F6B7A"
BLUE = "#2E74B5"
DARK_BLUE = "#1F4D78"
GREEN = "#167A5B"
GOLD = "#9A6B00"
RED = "#B53D44"
PURPLE = "#6B5BAA"
TEAL = "#007C89"
GRAY_LINE = "#D7DEE8"
SOFT_BG = "#F7F9FC"


def ensure_dirs() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    PDF_PATH.parent.mkdir(parents=True, exist_ok=True)
    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(FONT_BOLD_PATH if bold else FONT_PATH), size)


def draw_round_rect(draw: ImageDraw.ImageDraw, box, fill, outline=None, radius=24, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def text_width(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont) -> float:
    return draw.textlength(text, font=fnt)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for para in text.split("\n"):
        current = ""
        for ch in para:
            candidate = current + ch
            if not current or text_width(draw, candidate, fnt) <= max_width:
                current = candidate
            else:
                lines.append(current)
                current = ch
        if current:
            lines.append(current)
    return lines


def draw_wrapped(
    draw: ImageDraw.ImageDraw,
    text: str,
    xy: tuple[int, int],
    max_width: int,
    fnt: ImageFont.FreeTypeFont,
    fill=INK,
    line_gap: int = 8,
    max_lines: int | None = None,
) -> int:
    x, y = xy
    lines = wrap_text(draw, text, fnt, max_width)
    if max_lines is not None:
        lines = lines[:max_lines]
    line_h = fnt.size + line_gap
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=fill)
        y += line_h
    return y


def draw_arrow(draw: ImageDraw.ImageDraw, start, end, fill="#6B7480", width=5):
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=fill, width=width)
    angle = 0 if x2 >= x1 else pi
    if abs(y2 - y1) > abs(x2 - x1):
        angle = pi / 2 if y2 >= y1 else -pi / 2
    size = 18
    pts = [
        (x2, y2),
        (x2 - size * cos(angle - 0.45), y2 - size * sin(angle - 0.45)),
        (x2 - size * cos(angle + 0.45), y2 - size * sin(angle + 0.45)),
    ]
    draw.polygon(pts, fill=fill)


def make_continuum(path: Path) -> None:
    img = Image.new("RGB", (1800, 740), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1800, 740), fill="#FFFFFF")
    title = font(44, True)
    body = font(29)
    small = font(23)
    draw.text((70, 52), "核心地图：从平庸到卓越", font=title, fill=INK)
    draw.text((70, 112), "不是否定自回归，而是判断概率路径与任务价值是否同向。", font=small, fill=MUTED)

    cards = [
        ("LLM 平庸", "流畅、合理、可迭代\n但停在低价值盆地", "#FCEBED", RED),
        ("局部对齐", "局部步骤有价值\n全局仍可能偏移", "#FFF4D8", GOLD),
        ("LLM 卓越", "局部续写与全局价值\n彼此增强", "#E8F6EE", GREEN),
    ]
    x0, y0, w, h = 80, 230, 500, 280
    for i, (name, desc, bg, accent) in enumerate(cards):
        x = x0 + i * 610
        draw_round_rect(draw, (x, y0, x + w, y0 + h), fill=bg, outline=accent, radius=34, width=4)
        draw.ellipse((x + 34, y0 + 34, x + 112, y0 + 112), fill=accent)
        draw.text((x + 58, y0 + 52), str(i + 1), font=font(34, True), fill="white")
        draw.text((x + 140, y0 + 44), name, font=font(42, True), fill=accent)
        draw_wrapped(draw, desc, (x + 48, y0 + 145), w - 96, body, fill=INK, line_gap=14)
        if i < 2:
            draw_arrow(draw, (x + w + 25, y0 + h // 2), (x + w + 95, y0 + h // 2), fill="#8A96A3")

    draw_round_rect(draw, (170, 580, 1630, 665), fill=SOFT_BG, outline=GRAY_LINE, radius=24, width=2)
    draw.text((210, 606), "干预原则", font=font(30, True), fill=DARK_BLUE)
    draw.text(
        (365, 606),
        "保留已对齐的部分，把高失配任务改造成更低失配、更接近卓越的子任务。",
        font=font(28),
        fill=INK,
    )
    img.save(path)


def make_mismatch_matrix(path: Path) -> None:
    img = Image.new("RGB", (1800, 1120), "white")
    draw = ImageDraw.Draw(img)
    draw.text((70, 48), "六类原始失配：为什么搜索会停在“还不错”", font=font(42, True), fill=INK)
    draw.text((70, 108), "它们不是表面症状，而是可达生成与真实任务价值分岔的诊断轴。", font=font(24), fill=MUTED)

    items = [
        ("聚合失配", "局部合理的续写，无法组合成全局高价值结构。", RED, "治：建立结构大纲与全局审计。"),
        ("支持失配", "高价值结构处在低概率或低支持区域，概率无法区分洞见与噪声。", PURPLE, "治：显式枚举稀有结构与反例。"),
        ("状态失配", "效用依赖当前观测通道下不可识别的潜在状态。", TEAL, "治：补观测、建状态表、设置升级规则。"),
        ("规格失配", "提示词、训练规范或评估器偏离真实任务标准。", GOLD, "治：先生成评分规约，再产出结果。"),
        ("拟合边界失配", "已学能力的触发边界与真实适用边界不一致。", GREEN, "治：路由条件、禁用条件、触发测试。"),
        ("观测-表征失配", "观测、编码、上下文或工具通道丢失了任务充分变量。", BLUE, "治：改变表征与控制空间。"),
    ]
    card_w, card_h = 520, 260
    start_x, start_y = 70, 210
    gap_x, gap_y = 50, 46
    for idx, (name, desc, accent, fix) in enumerate(items):
        row, col = divmod(idx, 3)
        x = start_x + col * (card_w + gap_x)
        y = start_y + row * (card_h + gap_y)
        draw_round_rect(draw, (x, y, x + card_w, y + card_h), fill="#FFFFFF", outline=GRAY_LINE, radius=24, width=2)
        draw.rectangle((x, y, x + 18, y + card_h), fill=accent)
        draw.text((x + 42, y + 34), name, font=font(34, True), fill=accent)
        draw_wrapped(draw, desc, (x + 42, y + 91), card_w - 82, font(25), fill=INK, line_gap=9)
        draw_wrapped(draw, fix, (x + 42, y + 185), card_w - 82, font(22), fill=MUTED, line_gap=8)

    draw_round_rect(draw, (110, 850, 1690, 1000), fill="#F7FBFF", outline="#C9D8EA", radius=28, width=2)
    draw.text((160, 890), "读法", font=font(32, True), fill=DARK_BLUE)
    draw_wrapped(
        draw,
        "先定位主要失配，再选择对应的控制动作。很多失败现象只是这些原始失配在预算、表征和控制策略下的复合结果。",
        (260, 888),
        1340,
        font(27),
        fill=INK,
        line_gap=10,
    )
    img.save(path)


def make_transformation_flow(path: Path) -> None:
    img = Image.new("RGB", (1800, 980), "white")
    draw = ImageDraw.Draw(img)
    draw.text((70, 50), "从平庸到卓越：把任务换一种呈现方式", font=font(42, True), fill=INK)
    draw.text((70, 108), "关键不是让模型停止生成，而是把困难目标拆成更容易正向对齐的中间任务。", font=font(24), fill=MUTED)

    nodes = [
        ("困难最终输出", "目标模糊、状态隐藏、全局约束强", "#FCEBED", RED),
        ("解耦控制空间", "先建模、枚举、约束、评分", "#FFF4D8", GOLD),
        ("受治理知识", "可验证、可撤销、可复用", "#E8F6EE", GREEN),
        ("流畅表达", "让自回归生成承担它擅长的部分", "#EAF3FB", BLUE),
    ]
    x_positions = [90, 520, 960, 1390]
    y, w, h = 250, 330, 220
    for i, (title, desc, bg, accent) in enumerate(nodes):
        x = x_positions[i]
        draw_round_rect(draw, (x, y, x + w, y + h), fill=bg, outline=accent, radius=28, width=3)
        draw.text((x + 28, y + 34), title, font=font(31, True), fill=accent)
        draw_wrapped(draw, desc, (x + 28, y + 94), w - 56, font(24), fill=INK, line_gap=9)
        if i < len(nodes) - 1:
            draw_arrow(draw, (x + w + 18, y + h // 2), (x + w + 86, y + h // 2), fill="#7E8A97")

    draw.text((90, 570), "可复用操作模式", font=font(34, True), fill=INK)
    chips = [
        ("压缩上下文", BLUE),
        ("提取决策变量", TEAL),
        ("生成评分规约", GOLD),
        ("枚举边界情形", RED),
        ("构造查询", PURPLE),
        ("建立结构大纲", GREEN),
        ("诱导约束", DARK_BLUE),
        ("语义解压", "#7A5A00"),
    ]
    x, y2 = 90, 645
    for label, accent in chips:
        tw = int(text_width(draw, label, font(25, True))) + 58
        if x + tw > 1690:
            x = 90
            y2 += 84
        draw_round_rect(draw, (x, y2, x + tw, y2 + 58), fill="#FFFFFF", outline=accent, radius=18, width=3)
        draw.text((x + 29, y2 + 14), label, font=font(25, True), fill=accent)
        x += tw + 24

    draw_round_rect(draw, (90, 830, 1690, 925), fill=SOFT_BG, outline=GRAY_LINE, radius=26, width=2)
    draw.text((130, 858), "落地句式", font=font(28, True), fill=DARK_BLUE)
    draw_wrapped(
        draw,
        "不要直接问“给我最佳答案”。先问“需要哪些状态、标准、约束和反例，才能判定答案是高价值的？”",
        (270, 858),
        1330,
        font(25),
        fill=INK,
        line_gap=8,
    )
    img.save(path)


def make_governance_loop(path: Path) -> None:
    img = Image.new("RGB", (1800, 1040), "white")
    draw = ImageDraw.Draw(img)
    draw.text((70, 48), "知识治理：把控制知识从上下文叙事中拿出来", font=font(42, True), fill=INK)
    draw.text((70, 108), "最终表述与知识获取、验证、控制解耦，减少“看似合理”的漂移。", font=font(24), fill=MUTED)

    cx, cy, r = 610, 555, 300
    steps = [
        ("构造控制空间", -90, BLUE),
        ("提取候选不变项", -18, TEAL),
        ("证据验证", 54, GREEN),
        ("写入 GKO", 126, GOLD),
        ("引导生成", 198, PURPLE),
        ("监测失败并撤销", 270, RED),
    ]
    for label, deg, accent in steps:
        angle = deg * pi / 180
        x = cx + int(r * cos(angle))
        y = cy + int(r * sin(angle))
        draw.ellipse((x - 84, y - 84, x + 84, y + 84), fill="#FFFFFF", outline=accent, width=4)
        draw_wrapped(draw, label, (x - 62, y - 34), 124, font(23, True), fill=accent, line_gap=4)
    for _, deg, _ in steps:
        a1 = (deg + 24) * pi / 180
        a2 = (deg + 48) * pi / 180
        draw_arrow(
            draw,
            (cx + int((r - 10) * cos(a1)), cy + int((r - 10) * sin(a1))),
            (cx + int((r - 10) * cos(a2)), cy + int((r - 10) * sin(a2))),
            fill="#A3ACB8",
            width=4,
        )
    draw.ellipse((cx - 120, cy - 120, cx + 120, cy + 120), fill="#F7FBFF", outline="#C9D8EA", width=3)
    draw_wrapped(draw, "治理循环\n本身与任务无关", (cx - 78, cy - 42), 160, font(27, True), fill=DARK_BLUE, line_gap=8)

    card_x, card_y = 1010, 240
    draw_round_rect(draw, (card_x, card_y, 1690, 830), fill="#FFFFFF", outline=GRAY_LINE, radius=32, width=2)
    draw.rectangle((card_x, card_y, card_x + 680, card_y + 82), fill="#E8EEF5")
    draw.text((card_x + 40, card_y + 24), "GKO：受治理知识对象", font=font(31, True), fill=DARK_BLUE)
    fields = [
        ("内容", "一条任务特定的控制知识、约束或路由规则"),
        ("证据", "支持它的观测、测试、反例与来源"),
        ("适用条件", "何时启用，何时不得启用"),
        ("强度与优先级", "软偏好、硬约束、诊断测试或路由规则"),
        ("生命周期", "创建、复核、过期、撤销与冲突处理"),
    ]
    y = card_y + 120
    for name, desc in fields:
        draw.text((card_x + 44, y), name, font=font(24, True), fill=GREEN if name == "证据" else DARK_BLUE)
        draw_wrapped(draw, desc, (card_x + 185, y), 440, font(23), fill=INK, line_gap=7)
        y += 82

    draw_round_rect(draw, (card_x + 40, card_y + 480, card_x + 640, card_y + 545), fill="#FFF8E8", outline="#E0C26F", radius=20, width=2)
    draw.text((card_x + 72, card_y + 496), "SGAR：长程 agent 中的状态权威层", font=font(23, True), fill=GOLD)
    img.save(path)


def make_decision_map(path: Path) -> None:
    img = Image.new("RGB", (1800, 1040), "white")
    draw = ImageDraw.Draw(img)
    draw.text((70, 48), "何时需要知识治理：一个快速判断", font=font(42, True), fill=INK)
    draw.text((70, 108), "治理不是普适必要项。只有当局部生成优势不足以支配全局价值时，它才值得引入。", font=font(24), fill=MUTED)

    x0, y0, w, h = 230, 230, 1120, 640
    draw.rectangle((x0, y0, x0 + w, y0 + h), outline=GRAY_LINE, width=4)
    draw.line((x0 + w // 2, y0, x0 + w // 2, y0 + h), fill=GRAY_LINE, width=4)
    draw.line((x0, y0 + h // 2, x0 + w, y0 + h // 2), fill=GRAY_LINE, width=4)
    draw.text((x0 + 250, y0 + h + 35), "局部对齐程度低", font=font(24), fill=MUTED)
    draw.text((x0 + 740, y0 + h + 35), "局部对齐程度高", font=font(24), fill=MUTED)
    draw.text((70, y0 + 115), "失配强", font=font(24), fill=MUTED)
    draw.text((70, y0 + 455), "失配弱", font=font(24), fill=MUTED)

    quadrants = [
        (x0 + 28, y0 + 32, "先补表征", "缺状态、缺标准、缺观测时，先不要急着生成最终答案。", RED),
        (x0 + w // 2 + 28, y0 + 32, "用知识治理", "模型能生成好部件，但高价值行为依赖可外化、可验证的控制知识。", GREEN),
        (x0 + 28, y0 + h // 2 + 32, "重写任务", "先拆解、澄清、换接口，把任务改到可处理区间。", GOLD),
        (x0 + w // 2 + 28, y0 + h // 2 + 32, "直接生成即可", "上下文压缩、语体迁移、结构化转换等原生卓越任务，可少用治理。", BLUE),
    ]
    for x, y, title, desc, accent in quadrants:
        draw.text((x, y), title, font=font(32, True), fill=accent)
        draw_wrapped(draw, desc, (x, y + 60), 460, font(24), fill=INK, line_gap=9)

    draw_round_rect(draw, (1415, 285, 1695, 815), fill=SOFT_BG, outline=GRAY_LINE, radius=26, width=2)
    draw.text((1450, 320), "五问检查", font=font(28, True), fill=DARK_BLUE)
    checks = [
        "任务价值是否依赖隐藏状态？",
        "是否存在低概率但高价值结构？",
        "同类错误是否会跨候选重复出现？",
        "评分标准是否只能在看到候选后才清楚？",
        "需要复用的约束是否只停留在上下文里？",
    ]
    yy = 375
    for item in checks:
        draw.ellipse((1454, yy + 8, 1474, yy + 28), fill=GREEN)
        draw_wrapped(draw, item, (1490, yy), 170, font(22), fill=INK, line_gap=6)
        yy += 82
    img.save(path)


def build_assets() -> dict[str, Path]:
    assets = {
        "continuum": ASSET_DIR / "01_continuum.png",
        "mismatch": ASSET_DIR / "02_mismatch_matrix.png",
        "flow": ASSET_DIR / "03_transformation_flow.png",
        "governance": ASSET_DIR / "04_governance_loop.png",
        "decision": ASSET_DIR / "05_decision_map.png",
    }
    make_continuum(assets["continuum"])
    make_mismatch_matrix(assets["mismatch"])
    make_transformation_flow(assets["flow"])
    make_governance_loop(assets["governance"])
    make_decision_map(assets["decision"])
    return assets


def register_pdf_font() -> str:
    pdfmetrics.registerFont(TTFont("SimHei", str(SIMHEI)))
    return "SimHei"


def pdf_wrap(c: canvas.Canvas, text: str, font_name: str, size: int, max_width: float) -> list[str]:
    lines: list[str] = []
    for para in text.split("\n"):
        current = ""
        for ch in para:
            candidate = current + ch
            if not current or pdfmetrics.stringWidth(candidate, font_name, size) <= max_width:
                current = candidate
            else:
                lines.append(current)
                current = ch
        if current:
            lines.append(current)
    return lines


def pdf_text(
    c: canvas.Canvas,
    text: str,
    x: float,
    y: float,
    max_width: float,
    font_name: str,
    size: int,
    leading: int | None = None,
    color=colors.HexColor(INK),
) -> float:
    if leading is None:
        leading = int(size * 1.45)
    c.setFont(font_name, size)
    c.setFillColor(color)
    for line in pdf_wrap(c, text, font_name, size, max_width):
        c.drawString(x, y, line)
        y -= leading
    return y


def pdf_header(c: canvas.Canvas, title: str, page_no: int, font_name: str) -> None:
    width, height = letter
    c.setStrokeColor(colors.HexColor("#D9E0EA"))
    c.setLineWidth(0.8)
    c.line(54, height - 44, width - 54, height - 44)
    c.setFont(font_name, 9)
    c.setFillColor(colors.HexColor(MUTED))
    c.drawString(54, height - 34, "知识治理短版图解")
    c.drawRightString(width - 54, height - 34, f"{page_no}")
    c.setFont(font_name, 20)
    c.setFillColor(colors.HexColor(DARK_BLUE))
    c.drawString(54, height - 82, title)


def draw_pdf_image(c: canvas.Canvas, path: Path, x: float, y_top: float, width: float, height: float) -> None:
    c.drawImage(ImageReader(str(path)), x, y_top - height, width=width, height=height, preserveAspectRatio=True, mask="auto")


def build_pdf(assets: dict[str, Path]) -> None:
    font_name = register_pdf_font()
    c = canvas.Canvas(str(PDF_PATH), pagesize=letter)
    width, height = letter
    margin = 54

    # Page 1
    c.setFillColor(colors.HexColor("#F7FBFF"))
    c.rect(0, 0, width, height, fill=1, stroke=0)
    c.setFillColor(colors.HexColor(DARK_BLUE))
    c.rect(0, height - 12, width, 12, fill=1, stroke=0)
    c.setFillColor(colors.HexColor(GREEN))
    c.rect(0, height - 12, width * 0.28, 12, fill=1, stroke=0)
    c.setFillColor(colors.HexColor(GOLD))
    c.rect(width * 0.28, height - 12, width * 0.12, 12, fill=1, stroke=0)
    c.setFillColor(colors.HexColor("#FFFFFF"))
    c.roundRect(margin, height - 248, width - 2 * margin, 94, 12, fill=1, stroke=0)
    c.setFont(font_name, 9)
    c.setFillColor(colors.HexColor(MUTED))
    c.drawString(margin + 18, height - 178, SOURCE_URL)
    c.setStrokeColor(colors.HexColor("#E4EAF2"))
    c.setLineWidth(0.8)
    c.line(margin + 18, height - 194, width - margin - 18, height - 194)
    c.setFillColor(colors.HexColor(DARK_BLUE))
    c.setFont(font_name, 30)
    c.drawString(margin, height - 88, "大语言模型系统的知识治理")
    c.setFont(font_name, 18)
    c.setFillColor(colors.HexColor(INK))
    c.drawString(margin, height - 124, "从 LLM 平庸，经由局部对齐，走向 LLM 卓越 - 短版图解")
    pdf_text(
        c,
        "一句话：LLM 的关键问题不只是“会不会生成”，而是它观察、表征、指定、支持与聚合的东西，是否和真实任务价值对齐。",
        margin + 18,
        height - 213,
        width - 2 * margin - 36,
        font_name,
        12,
        17,
        colors.HexColor(INK),
    )
    draw_pdf_image(c, assets["continuum"], margin, height - 270, width - 2 * margin, 225)
    y = 258
    c.setFont(font_name, 16)
    c.setFillColor(colors.HexColor(DARK_BLUE))
    c.drawString(margin, y, "三个保留点")
    y -= 30
    takeaways = [
        "LLM 平庸不是模型的普遍属性，而是任务、表征、状态、规格和预算共同形成的运行区间。",
        "局部对齐是常态：模型能生成有价值部件，但全局目标仍可能漂移。",
        "知识治理的目的，是把可验证、可撤销、可复用的控制知识从上下文叙事中外化出来。",
    ]
    for item in takeaways:
        c.setFillColor(colors.HexColor(GREEN))
        c.circle(margin + 5, y + 4, 3, fill=1, stroke=0)
        y = pdf_text(c, item, margin + 18, y + 8, width - 2 * margin - 18, font_name, 11, 16)
        y -= 4
    c.setFont(font_name, 8)
    c.setFillColor(colors.HexColor(MUTED))
    c.drawString(margin, 34, "知识治理短版图解")
    c.showPage()

    # Page 2
    pdf_header(c, "1. 六类原始失配", 2, font_name)
    draw_pdf_image(c, assets["mismatch"], margin, height - 110, width - 2 * margin, 360)
    y = 290
    body = (
        "这六类失配的价值在于诊断。它们帮助系统判断：继续采样、批判、重排或润色是否仍然有效，"
        "还是已经需要改变观测、表征、任务规格或控制结构。"
    )
    y = pdf_text(c, body, margin, y, width - 2 * margin, font_name, 12, 18)
    y -= 18
    pdf_text(
        c,
        "派生失败模式，例如顺序敏感轨迹、语料先验主导、预算下的控制容量崩塌，通常不是新的原始类别，而是六类失配叠加预算与控制策略后的表现。",
        margin,
        y,
        width - 2 * margin,
        font_name,
        11,
        17,
        colors.HexColor(MUTED),
    )
    c.showPage()

    # Page 3
    pdf_header(c, "2. 从平庸到卓越的转化", 3, font_name)
    draw_pdf_image(c, assets["flow"], margin, height - 110, width - 2 * margin, 315)
    y = 340
    pdf_text(
        c,
        "转化的实质，是把一个高失配的最终输出问题，改写成多个更接近 LLM 卓越区间的子任务。"
        "例如先压缩上下文、提取变量、生成评分规约、枚举边界情形，再让模型写最终答案。",
        margin,
        y,
        width - 2 * margin,
        font_name,
        12,
        18,
    )
    y -= 92
    c.setFillColor(colors.HexColor("#F7FBFF"))
    c.roundRect(margin, y - 60, width - 2 * margin, 82, 10, fill=1, stroke=0)
    pdf_text(
        c,
        "判断标准：如果子任务本身更容易被模型稳定完成，并且能为后续生成提供可检查的控制条件，那么转化就是有价值的。",
        margin + 16,
        y - 5,
        width - 2 * margin - 32,
        font_name,
        12,
        18,
        colors.HexColor(DARK_BLUE),
    )
    c.showPage()

    # Page 4
    pdf_header(c, "3. 知识治理与 GKO", 4, font_name)
    draw_pdf_image(c, assets["governance"], margin, height - 110, width - 2 * margin, 340)
    y = 315
    pdf_text(
        c,
        "知识治理不是把所有过程都写进提示词，而是创建一个被承认的控制层。候选知识先被验证，再以 GKO 形式记录条件、证据、强度、优先级、生命周期与撤销规则。",
        margin,
        y,
        width - 2 * margin,
        font_name,
        12,
        18,
    )
    y -= 92
    pdf_text(
        c,
        "对长程 agent 而言，SGAR 的重点是状态权威：计划、行动、观测、验证、升级、审计发现与撤销，不再只依赖上下文叙事，而成为可检查的状态转移。",
        margin,
        y,
        width - 2 * margin,
        font_name,
        11,
        17,
        colors.HexColor(MUTED),
    )
    c.showPage()

    # Page 5
    pdf_header(c, "4. 何时用，如何落地", 5, font_name)
    draw_pdf_image(c, assets["decision"], margin, height - 110, width - 2 * margin, 335)
    y = 318
    c.setFont(font_name, 15)
    c.setFillColor(colors.HexColor(DARK_BLUE))
    c.drawString(margin, y, "最小落地清单")
    y -= 28
    checklist = [
        "先标注任务中的已对齐部分：哪些生成、压缩、转换或枚举可以直接交给模型。",
        "再标注高失配部分：隐藏状态、低支持结构、规格模糊、全局聚合或表征损失。",
        "为高失配部分生成控制工件：评分规约、状态表、反例集、边界条件、路由规则。",
        "把验证后的控制工件写成 GKO，并设置何时启用、何时撤销。",
        "最后让模型只在受治理的边界内生成流畅表达，并用同一组 GKO 审计结果。",
    ]
    for item in checklist:
        c.setFillColor(colors.HexColor(GREEN))
        c.circle(margin + 5, y + 4, 3, fill=1, stroke=0)
        y = pdf_text(c, item, margin + 18, y + 8, width - 2 * margin - 18, font_name, 11, 16)
        y -= 3
    y -= 10
    c.setFillColor(colors.HexColor("#FFF8E8"))
    c.roundRect(margin, y - 54, width - 2 * margin, 72, 10, fill=1, stroke=0)
    pdf_text(
        c,
        "反面边界：如果任务本身已经处在 LLM 卓越区间，例如纯上下文压缩、语体迁移、结构化转换，普通提示或有限搜索可能更高效。",
        margin + 16,
        y - 4,
        width - 2 * margin - 32,
        font_name,
        11,
        16,
        colors.HexColor(GOLD),
    )
    c.showPage()
    c.save()


def set_run_font(run, size=None, color=None, bold=None, italic=None, east_asia="Microsoft YaHei"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size=None, color=None, bold=None, east_asia="Microsoft YaHei"):
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.allow_autofit = False
    table.columns[0].width = Inches(6.35)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7FBFF")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=11, color=DARK_BLUE.replace("#", ""), bold=True)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13, color=(BLUE if level in (1, 2) else DARK_BLUE).replace("#", ""), bold=True)


def add_para(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, size=11, bold=True, color=INK.replace("#", ""))
        r2 = p.add_run(text[len(bold_lead) :])
        set_run_font(r2, size=11, color=INK.replace("#", ""))
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK.replace("#", ""))


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK.replace("#", ""))


def add_figure(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.2))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(8)
    r = cp.add_run(caption)
    set_run_font(r, size=9, color=MUTED.replace("#", ""), italic=True)


def setup_docx_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=11, color=INK.replace("#", ""))
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[name]
        set_style_font(style, size=size, color=color.replace("#", ""), bold=True)
        style.paragraph_format.space_before = Pt(18 if name == "Heading 1" else 14)
        style.paragraph_format.space_after = Pt(10 if name == "Heading 1" else 7)
    set_style_font(styles["List Bullet"], size=11, color=INK.replace("#", ""))

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("知识治理短版图解")
    set_run_font(hr, size=9, color=MUTED.replace("#", ""))
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("知识治理短版图解")
    set_run_font(fr, size=8, color=MUTED.replace("#", ""))


def build_docx(assets: dict[str, Path]) -> None:
    doc = Document()
    setup_docx_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    tr = title.add_run("大语言模型系统的知识治理")
    set_run_font(tr, size=24, bold=True, color=DARK_BLUE.replace("#", ""))

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    sr = subtitle.add_run("从 LLM 平庸，经由局部对齐，走向 LLM 卓越 - 短版图解")
    set_run_font(sr, size=13, color=MUTED.replace("#", ""))

    meta = doc.add_paragraph()
    mr = meta.add_run(SOURCE_URL)
    set_run_font(mr, size=9.5, color=MUTED.replace("#", ""))

    add_callout(
        doc,
        "一句话：LLM 的关键问题不只是“会不会生成”，而是它观察、表征、指定、支持与聚合的东西，是否和真实任务价值对齐。",
    )
    add_figure(doc, assets["continuum"], "图 1：三段式视角 - 平庸、局部对齐与卓越")
    add_heading(doc, "三个保留点", level=1)
    for item in [
        "LLM 平庸是一种依赖任务、表征、状态、规格和预算的运行区间。",
        "局部对齐是常态：模型能生成有价值部件，但全局目标仍可能漂移。",
        "知识治理把可验证、可撤销、可复用的控制知识从上下文叙事中外化出来。",
    ]:
        add_bullet(doc, item)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "1. 六类原始失配", level=1)
    add_figure(doc, assets["mismatch"], "图 2：六类原始失配与对应治理动作")
    add_para(
        doc,
        "六类失配的价值在于诊断。它们帮助系统判断：继续采样、批判、重排或润色是否仍然有效，还是已经需要改变观测、表征、任务规格或控制结构。",
    )
    add_para(
        doc,
        "派生失败模式，例如顺序敏感轨迹、语料先验主导、预算下的控制容量崩塌，通常不是新的原始类别，而是六类失配叠加预算与控制策略后的表现。",
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "2. 从平庸到卓越的转化", level=1)
    add_figure(doc, assets["flow"], "图 3：把困难最终输出改造成更低失配的中间任务")
    add_para(
        doc,
        "转化的实质，是把一个高失配的最终输出问题，改写成多个更接近 LLM 卓越区间的子任务。典型顺序是：压缩上下文，提取变量，生成评分规约，枚举边界情形，再让模型写最终答案。",
    )
    add_callout(
        doc,
        "判断标准：如果子任务本身更容易被模型稳定完成，并且能为后续生成提供可检查的控制条件，那么转化就是有价值的。",
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "3. 知识治理与 GKO", level=1)
    add_figure(doc, assets["governance"], "图 4：治理循环与受治理知识对象 GKO")
    add_para(
        doc,
        "知识治理不是把所有过程都写进提示词，而是创建一个被承认的控制层。候选知识先被验证，再以 GKO 形式记录条件、证据、强度、优先级、生命周期与撤销规则。",
    )
    add_para(
        doc,
        "对长程 agent 而言，SGAR 的重点是状态权威：计划、行动、观测、验证、升级、审计发现与撤销，不再只依赖上下文叙事，而成为可检查的状态转移。",
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "4. 何时用，如何落地", level=1)
    add_figure(doc, assets["decision"], "图 5：何时需要知识治理的快速判断")
    add_heading(doc, "最小落地清单", level=2)
    for item in [
        "先标注任务中的已对齐部分：哪些生成、压缩、转换或枚举可以直接交给模型。",
        "再标注高失配部分：隐藏状态、低支持结构、规格模糊、全局聚合或表征损失。",
        "为高失配部分生成控制工件：评分规约、状态表、反例集、边界条件、路由规则。",
        "把验证后的控制工件写成 GKO，并设置何时启用、何时撤销。",
        "最后让模型只在受治理的边界内生成流畅表达，并用同一组 GKO 审计结果。",
    ]:
        add_bullet(doc, item)
    add_callout(
        doc,
        "反面边界：如果任务本身已经处在 LLM 卓越区间，例如纯上下文压缩、语体迁移、结构化转换，普通提示或有限搜索可能更高效。",
    )
    doc.save(DOCX_PATH)


def main() -> None:
    ensure_dirs()
    assets = build_assets()
    build_pdf(assets)
    build_docx(assets)
    print(PDF_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
